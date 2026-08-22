"""Services métier du module 9 : tableau de bord et reporting.

Fournit les KPIs temps réel et graphiques du dashboard principal, les neuf
rapports prédéfinis, le moteur de rapports personnalisés (datasets, filtres
avancés, groupements, planification, partage), les exports PDF / Excel /
CSV / Word / JSON et les alertes paramétrables à seuils personnalisés.
"""

import csv
import io
from datetime import date, datetime, timedelta, timezone
from typing import Any, Dict, List, Optional, Tuple

from sqlalchemy import Column as SaColumn
from sqlalchemy import func as sa_func
from sqlalchemy.orm import Session

from app.models.crm import (
    DealStatus,
    Listing,
    PipelineDeal,
    Prospect,
    ProspectStatus,
    Visit,
    VisitStatus,
)
from app.models.finance import Charge, ChargeRecoverability, LatePayment
from app.models.maintenance import MaintenanceTicket, TicketStatus, WorkProject
from app.models.owner import Mandate, Owner
from app.models.property import Property
from app.models.reporting import (
    AlertEvent,
    AlertRule,
    CustomReport,
    DashboardWidget,
    ReportExecution,
    ScheduleFrequency,
)
from app.models.tenant import Lease, LeaseStatus, RentPayment


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _as_utc(value: Optional[datetime]) -> Optional[datetime]:
    if value is None or value.tzinfo is None:
        return value.replace(tzinfo=timezone.utc) if value else None
    return value


def _month_start(d: date) -> date:
    return d.replace(day=1)


def _months_back(reference: date, months: int) -> date:
    month = reference.month - months
    year = reference.year
    while month <= 0:
        month += 12
        year -= 1
    return date(year, month, 1)


# ---------------------------------------------------------------------------
# Dashboard principal : KPIs temps réel
# ---------------------------------------------------------------------------
def dashboard_kpis(db: Session) -> Dict:
    today = date.today()
    month_s, month_e = _month_start(today), today
    year_s = date(today.year, 1, 1)

    total_properties = db.query(Property).filter(Property.is_active == True).count()  # noqa: E712
    rentable = db.query(Property).filter(
        Property.is_active == True, Property.rent_price != None  # noqa: E712
    ).count()
    occupied = db.query(Lease).filter(Lease.status == LeaseStatus.ACTIVE).count()
    occupancy_rate = round(occupied / rentable * 100, 2) if rentable else 0.0

    def _collected(start: date, end: date) -> float:
        return (
            db.query(sa_func.coalesce(sa_func.sum(RentPayment.amount_paid), 0))
            .filter(
                RentPayment.paid_at >= datetime.combine(start, datetime.min.time()),
                RentPayment.paid_at <= datetime.combine(end, datetime.max.time()),
            )
            .scalar()
            or 0
        )

    monthly_revenue = _collected(month_s, month_e)
    annual_revenue = _collected(year_s, today)
    target_annual = monthly_revenue * 12

    unpaid_outstanding = (
        db.query(sa_func.coalesce(sa_func.sum(LatePayment.amount_outstanding), 0))
        .filter(LatePayment.status != "resolved")
        .scalar()
        or 0
    )
    unpaid_count = (
        db.query(LatePayment).filter(LatePayment.status != "resolved").count()
    )

    open_tickets = (
        db.query(MaintenanceTicket)
        .filter(
            MaintenanceTicket.status.notin_(
                [TicketStatus.COMPLETED, TicketStatus.CLOSED, TicketStatus.CANCELLED]
            )
        )
        .count()
    )

    def _expiring(days: int) -> int:
        return (
            db.query(Lease)
            .filter(
                Lease.status == LeaseStatus.ACTIVE,
                Lease.end_date >= today,
                Lease.end_date <= today + timedelta(days=days),
            )
            .count()
        )

    expiring_leases = {"30": _expiring(30), "60": _expiring(60), "90": _expiring(90)}
    expiring_mandates = (
        db.query(Mandate)
        .filter(
            Mandate.status == "active",
            Mandate.end_date >= today,
            Mandate.end_date <= today + timedelta(days=90),
        )
        .count()
    )

    active_prospects = (
        db.query(Prospect).filter(Prospect.status == ProspectStatus.ACTIVE).count()
    )
    open_deals_value = (
        db.query(sa_func.coalesce(sa_func.sum(PipelineDeal.estimated_value), 0))
        .filter(PipelineDeal.status == DealStatus.OPEN)
        .scalar()
        or 0
    )

    return {
        "generated_at": _now().isoformat(),
        "properties": {
            "total": total_properties,
            "rentable": rentable,
            "occupied": occupied,
            "vacant": max(rentable - occupied, 0),
        },
        "occupancy_rate_pct": occupancy_rate,
        "revenue": {
            "month_to_date": round(monthly_revenue, 2),
            "year_to_date": round(annual_revenue, 2),
            "annual_run_rate": round(target_annual, 2),
        },
        "unpaid": {"count": unpaid_count, "outstanding": round(unpaid_outstanding, 2)},
        "maintenance": {"open_tickets": open_tickets},
        "leases_expiring": expiring_leases,
        "mandates_expiring_90d": expiring_mandates,
        "crm": {
            "active_prospects": active_prospects,
            "open_deals": db.query(PipelineDeal).filter(PipelineDeal.status == DealStatus.OPEN).count(),
            "open_deals_value": round(open_deals_value, 2),
            "visits_upcoming": db.query(Visit)
            .filter(
                Visit.scheduled_date >= today,
                Visit.status.in_([VisitStatus.SCHEDULED, VisitStatus.CONFIRMED]),
            )
            .count(),
        },
    }


# ---------------------------------------------------------------------------
# Dashboard : graphiques dynamiques
# ---------------------------------------------------------------------------
def dashboard_charts(db: Session, months: int = 12) -> Dict:
    today = date.today()

    # Évolution des revenus (encaissements par mois)
    revenue_series = []
    for i in range(months - 1, -1, -1):
        start = _months_back(today, i)
        if start.month == 12:
            end = date(start.year + 1, 1, 1) - timedelta(days=1)
        else:
            end = date(start.year, start.month + 1, 1) - timedelta(days=1)
        end = min(end, today)
        amount = (
            db.query(sa_func.coalesce(sa_func.sum(RentPayment.amount_paid), 0))
            .filter(
                RentPayment.paid_at >= datetime.combine(start, datetime.min.time()),
                RentPayment.paid_at <= datetime.combine(end, datetime.max.time()),
            )
            .scalar()
            or 0
        )
        revenue_series.append({"month": start.isoformat(), "amount": round(amount, 2)})

    # Répartition par type de bien
    type_rows = (
        db.query(Property.type, sa_func.count(Property.id))
        .filter(Property.is_active == True)  # noqa: E712
        .group_by(Property.type)
        .all()
    )
    property_types = [
        {"type": (t.value if t else "autre"), "count": count} for t, count in type_rows
    ]

    # Taux d'occupation mensuel : baux actifs par rapport au parc louable
    occupancy_series = []
    total_rentable = db.query(Property).filter(
        Property.is_active == True, Property.rent_price != None  # noqa: E712
    ).count()
    for i in range(months - 1, -1, -1):
        month = _months_back(today, i)
        if month.month == 12:
            month_end = date(month.year + 1, 1, 1) - timedelta(days=1)
        else:
            month_end = date(month.year, month.month + 1, 1) - timedelta(days=1)
        active = (
            db.query(Lease)
            .filter(
                Lease.start_date <= month_end,
                (Lease.end_date == None) | (Lease.end_date >= month),  # noqa: E711
                Lease.status.in_([LeaseStatus.ACTIVE, LeaseStatus.TERMINATED, LeaseStatus.EXPIRED]),
                Lease.start_date <= month_end,
            )
            .count()
        )
        rate = round(active / total_rentable * 100, 1) if total_rentable else 0.0
        occupancy_series.append({"month": month.isoformat(), "rate": rate})

    # Répartition des charges par nature
    charge_rows = (
        db.query(Charge.charge_type, sa_func.coalesce(sa_func.sum(Charge.amount), 0))
        .group_by(Charge.charge_type)
        .all()
    )
    charge_distribution = [
        {"type": ctype or "autre", "amount": round(amount, 2)} for ctype, amount in charge_rows
    ]

    # Performance commerciale : dossiers créés / gagnés par mois
    commercial_series = []
    for i in range(months - 1, -1, -1):
        month = _months_back(today, i)
        if month.month == 12:
            month_end = date(month.year + 1, 1, 1) - timedelta(days=1)
        else:
            month_end = date(month.year, month.month + 1, 1) - timedelta(days=1)
        created = (
            db.query(PipelineDeal)
            .filter(
                PipelineDeal.created_at >= datetime.combine(month, datetime.min.time()),
                PipelineDeal.created_at <= datetime.combine(month_end, datetime.max.time()),
            )
            .count()
        )
        won = (
            db.query(PipelineDeal)
            .filter(
                PipelineDeal.closed_at >= datetime.combine(month, datetime.min.time()),
                PipelineDeal.closed_at <= datetime.combine(month_end, datetime.max.time()),
                PipelineDeal.status == DealStatus.WON,
            )
            .count()
        )
        commercial_series.append(
            {"month": month.isoformat(), "created": created, "won": won}
        )

    return {
        "revenue_evolution": revenue_series,
        "property_type_distribution": property_types,
        "occupancy_monthly": occupancy_series,
        "charge_distribution": charge_distribution,
        "commercial_performance": commercial_series,
    }


# ---------------------------------------------------------------------------
# Widgets personnalisables (drag & drop)
# ---------------------------------------------------------------------------
WIDGET_CATALOG = [
    {"type": "kpi_occupancy", "title": "Taux d'occupation", "default_size": "small"},
    {"type": "kpi_revenue", "title": "Revenus du mois", "default_size": "small"},
    {"type": "kpi_unpaid", "title": "Impayés en cours", "default_size": "small"},
    {"type": "kpi_tickets", "title": "Tickets maintenance", "default_size": "small"},
    {"type": "kpi_leases_expiring", "title": "Baux à échéance", "default_size": "small"},
    {"type": "kpi_prospects", "title": "Prospects actifs", "default_size": "small"},
    {"type": "chart_revenue", "title": "Évolution des revenus", "default_size": "large"},
    {"type": "chart_occupancy", "title": "Taux d'occupation mensuel", "default_size": "medium"},
    {"type": "chart_property_types", "title": "Répartition par type de bien", "default_size": "medium"},
    {"type": "chart_charges", "title": "Répartition des charges", "default_size": "medium"},
    {"type": "chart_commercial", "title": "Performance commerciale", "default_size": "large"},
    {"type": "list_upcoming_visits", "title": "Prochaines visites", "default_size": "medium"},
    {"type": "list_open_tickets", "title": "Tickets ouverts", "default_size": "medium"},
    {"type": "list_latest_payments", "title": "Derniers encaissements", "default_size": "medium"},
]


def list_widgets(db: Session, user_email: str) -> List[DashboardWidget]:
    return (
        db.query(DashboardWidget)
        .filter(DashboardWidget.user_email == user_email)
        .order_by(DashboardWidget.column_index, DashboardWidget.position)
        .all()
)


def create_widget(db: Session, user_email: str, data) -> DashboardWidget:
    widget = DashboardWidget(user_email=user_email, **data.model_dump())
    db.add(widget)
    db.commit()
    db.refresh(widget)
    return widget


def update_widget(db: Session, widget_id: int, data) -> DashboardWidget:
    widget = db.query(DashboardWidget).filter(DashboardWidget.id == widget_id).first()
    if not widget:
        raise ValueError("Widget non trouvé")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(widget, field, value)
    db.commit()
    db.refresh(widget)
    return widget


def delete_widget(db: Session, widget_id: int) -> None:
    widget = db.query(DashboardWidget).filter(DashboardWidget.id == widget_id).first()
    if not widget:
        raise ValueError("Widget non trouvé")
    db.delete(widget)
    db.commit()


def reorder_widgets(db: Session, user_email: str, positions) -> List[DashboardWidget]:
    widgets = {
        w.id: w for w in db.query(DashboardWidget).filter(DashboardWidget.user_email == user_email).all()
    }
    for item in positions.positions:
        widget = widgets.get(item.widget_id)
        if not widget:
            raise ValueError(f"Widget {item.widget_id} non trouvé")
        widget.column_index = item.column_index
        widget.position = item.position
    db.commit()
    return list_widgets(db, user_email)


def widget_data(db: Session, widget_type: str) -> Dict:
    """Données temps réel associées à un type de widget."""
    kpis = dashboard_kpis(db)
    charts = dashboard_charts(db, months=6)
    mapping = {
        "kpi_occupancy": lambda: {"value": kpis["occupancy_rate_pct"], "unit": "%"},
        "kpi_revenue": lambda: {"value": kpis["revenue"]["month_to_date"], "unit": "€"},
        "kpi_unpaid": lambda: kpis["unpaid"],
        "kpi_tickets": lambda: kpis["maintenance"],
        "kpi_leases_expiring": lambda: kpis["leases_expiring"],
        "kpi_prospects": lambda: kpis["crm"],
        "chart_revenue": lambda: charts["revenue_evolution"],
        "chart_occupancy": lambda: charts["occupancy_monthly"],
        "chart_property_types": lambda: charts["property_type_distribution"],
        "chart_charges": lambda: charts["charge_distribution"],
        "chart_commercial": lambda: charts["commercial_performance"],
        "list_upcoming_visits": lambda: {
            "items": [
                {
                    "date": v.scheduled_date.isoformat(),
                    "time": v.start_time,
                    "property_id": v.property_id,
                }
                for v in db.query(Visit)
                .filter(Visit.scheduled_date >= date.today())
                .order_by(Visit.scheduled_date)
                .limit(10)
                .all()
            ]
        },
        "list_open_tickets": lambda: {
            "items": [
                {
                    "reference": t.reference,
                    "title": t.title,
                    "status": t.status.value if t.status else None,
                    "property_id": t.property_id,
                }
                for t in db.query(MaintenanceTicket)
                .filter(
                MaintenanceTicket.status.notin_(
                    [TicketStatus.COMPLETED, TicketStatus.CLOSED, TicketStatus.CANCELLED]
                )
            )
                .order_by(MaintenanceTicket.reported_at.desc())
                .limit(10)
                .all()
            ]
        },
        "list_latest_payments": lambda: {
            "items": [
                {
                    "period": p.period,
                    "amount": p.amount_paid,
                    "status": p.status.value if p.status else None,
                }
                for p in db.query(RentPayment)
                .order_by(RentPayment.paid_at.desc().nullslast(), RentPayment.id.desc())
                .limit(10)
                .all()
            ]
        },
    }
    builder = mapping.get(widget_type)
    if not builder:
        raise ValueError(f"Type de widget inconnu : {widget_type}")
    return builder()


# ---------------------------------------------------------------------------
# Rapports prédéfinis
# ---------------------------------------------------------------------------
PREDEFINED_REPORTS = {
    "rapport_gestion_locative": "Rapport de gestion locative",
    "etat_loyers": "État des loyers",
    "synthese_impayes": "Synthèse des impayés",
    "etat_travaux": "État des travaux",
    "rapport_vacance_locative": "Rapport de vacance locative",
    "bilan_proprietaire": "Bilan financier par propriétaire",
    "bilan_bien": "Bilan financier par bien",
    "rapport_activite_agence": "Rapport d'activité de l'agence",
    "rapport_fiscal_annuel": "Rapport fiscal annuel",
}


def _active_lease_for(db: Session, property_id: int) -> Optional[Lease]:
    return (
        db.query(Lease)
        .filter(Lease.property_id == property_id, Lease.status == LeaseStatus.ACTIVE)
        .first()
    )


def build_predefined_report(db: Session, report_key: str, params: Dict) -> Dict:
    year = int(params.get("year") or date.today().year)
    period_start = date(year, int(params.get("month") or 1), 1)
    if params.get("month"):
        next_month = period_start.month + 1
        next_year = period_start.year + (1 if next_month > 12 else 0)
        next_month = 1 if next_month > 12 else next_month
        period_end = date(next_year, next_month, 1) - timedelta(days=1)
    else:
        period_end = date(year, 12, 31)

    if report_key == "rapport_gestion_locative":
        rows = []
        properties = db.query(Property).filter(Property.is_active == True).all()  # noqa: E712
        for prop in properties:
            lease = _active_lease_for(db, prop.id)
            payments = (
                db.query(RentPayment)
                .filter(
                    RentPayment.lease_id == lease.id,
                    RentPayment.due_date >= period_start,
                    RentPayment.due_date <= period_end,
                )
                .all()
                if lease
                else []
            )
            rows.append(
                {
                    "bien": f"{prop.reference} — {prop.title}",
                    "ville": prop.city,
                    "loue": lease is not None,
                    "locataire": (
                        f"{lease.tenant.first_name} {lease.tenant.last_name}"
                        if lease and lease.tenant
                        else None
                    ),
                    "loyer_cc": (lease.monthly_rent + (lease.monthly_charges or 0)) if lease else None,
                    "echeances": len(payments),
                    "encaisse": round(sum(p.amount_paid or 0 for p in payments), 2),
                    "dû": round(sum(p.amount_due or 0 for p in payments), 2),
                }
            )
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "period": {"start": period_start.isoformat(), "end": period_end.isoformat()},
            "columns": ["bien", "ville", "loue", "locataire", "loyer_cc", "echeances", "encaisse", "dû"],
            "rows": rows,
        }

    if report_key == "etat_loyers":
        payments = (
            db.query(RentPayment)
            .filter(RentPayment.due_date >= period_start, RentPayment.due_date <= period_end)
            .order_by(RentPayment.due_date)
            .all()
        )
        rows = [
            {
                "echeance": p.due_date.isoformat(),
                "periode": p.period,
                "bien": p.lease.property.title if p.lease and p.lease.property else None,
                "locataire": (
                    f"{p.lease.tenant.first_name} {p.lease.tenant.last_name}"
                    if p.lease and p.lease.tenant
                    else None
                ),
                "dû": p.amount_due,
                "paye": p.amount_paid,
                "solde": round((p.amount_due or 0) - (p.amount_paid or 0), 2),
                "statut": p.status.value if p.status else None,
            }
            for p in payments
        ]
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "period": {"start": period_start.isoformat(), "end": period_end.isoformat()},
            "columns": ["echeance", "periode", "bien", "locataire", "dû", "paye", "solde", "statut"],
            "rows": rows,
            "totals": {
                "dû": round(sum(r["dû"] or 0 for r in rows), 2),
                "paye": round(sum(r["paye"] or 0 for r in rows), 2),
                "solde": round(sum(r["solde"] or 0 for r in rows), 2),
            },
        }

    if report_key == "synthese_impayes":
        late = db.query(LatePayment).filter(LatePayment.status != "resolved").all()
        rows = [
            {
                "reference": l.reference,
                "periode": l.period,
                "bien": l.property.title if l.property else None,
                "locataire": (
                    f"{l.tenant.first_name} {l.tenant.last_name}" if l.tenant else None
                ),
                "restant_dû": l.amount_outstanding,
                "penalites": l.penalty_amount,
                "jours_retard": l.overdue_days,
                "etape": l.stage.value if l.stage else None,
            }
            for l in late
        ]
        by_stage: Dict[str, float] = {}
        for row in rows:
            by_stage[row["etape"]] = round(by_stage.get(row["etape"], 0) + (row["restant_dû"] or 0), 2)
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "columns": ["reference", "periode", "bien", "locataire", "restant_dû", "penalites", "jours_retard", "etape"],
            "rows": rows,
            "totals": {
                "restant_dû": round(sum(r["restant_dû"] or 0 for r in rows), 2),
                "dossiers": len(rows),
            },
            "par_etape": by_stage,
        }

    if report_key == "etat_travaux":
        tickets = (
            db.query(MaintenanceTicket)
            .filter(MaintenanceTicket.reported_at >= datetime.combine(period_start, datetime.min.time()))
            .order_by(MaintenanceTicket.reported_at.desc())
            .all()
        )
        rows = [
            {
                "reference": t.reference,
                "titre": t.title,
                "bien": t.property.title if t.property else None,
                "urgence": t.urgency.value if t.urgency else None,
                "statut": t.status.value if t.status else None,
                "cout_estime": t.estimated_cost,
                "cout_final": t.final_cost,
                "ouvert_le": t.reported_at.date().isoformat() if t.reported_at else None,
            }
            for t in tickets
        ]
        projects = [
            {
                "nom": p.name,
                "bien": p.property.title if p.property else None,
                "statut": p.status.value if hasattr(p, "status") and p.status else None,
                "budget": getattr(p, "budget", None),
            }
            for p in db.query(WorkProject).all()
        ]
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "columns": ["reference", "titre", "bien", "urgence", "statut", "cout_estime", "cout_final", "ouvert_le"],
            "rows": rows,
            "totals": {
                "tickets": len(rows),
                "cout_final": round(sum(r["cout_final"] or 0 for r in rows), 2),
            },
            "projets_travaux": projects,
        }

    if report_key == "rapport_vacance_locative":
        today = date.today()
        rows = []
        properties = db.query(Property).filter(
            Property.is_active == True, Property.rent_price != None  # noqa: E712
        ).all()
        for prop in properties:
            if _active_lease_for(db, prop.id):
                continue
            last_lease = (
                db.query(Lease)
                .filter(Lease.property_id == prop.id)
                .order_by(Lease.end_date.desc().nullslast())
                .first()
            )
            vacant_since = None
            days_vacant = None
            if last_lease and last_lease.end_date:
                vacant_since = last_lease.end_date + timedelta(days=1)
                days_vacant = (today - vacant_since).days
            elif prop.created_at:
                vacant_since = prop.created_at.date()
                days_vacant = (today - vacant_since).days
            rows.append(
                {
                    "bien": f"{prop.reference} — {prop.title}",
                    "ville": prop.city,
                    "loyer_cible": prop.rent_price,
                    "vacant_depuis": vacant_since.isoformat() if vacant_since else None,
                    "jours_vacance": days_vacant,
                    "perte_estimee": round((prop.rent_price or 0) * (days_vacant or 0) / 30, 2),
                }
            )
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "columns": ["bien", "ville", "loyer_cible", "vacant_depuis", "jours_vacance", "perte_estimee"],
            "rows": rows,
            "totals": {
                "biens_vacants": len(rows),
                "perte_mensuelle": round(sum(r["loyer_cible"] or 0 for r in rows), 2),
                "perte_cumulee": round(sum(r["perte_estimee"] or 0 for r in rows), 2),
            },
        }

    if report_key == "bilan_proprietaire":
        owners = db.query(Owner).all()
        rows = []
        for owner in owners:
            properties = owner.properties
            property_ids = [p.id for p in properties]
            collected = (
                db.query(sa_func.coalesce(sa_func.sum(RentPayment.amount_paid), 0))
                .join(Lease, RentPayment.lease_id == Lease.id)
                .filter(
                    Lease.property_id.in_(property_ids or [-1]),
                    RentPayment.due_date >= period_start,
                    RentPayment.due_date <= period_end,
                )
                .scalar()
                or 0
            )
            charges_total = (
                db.query(sa_func.coalesce(sa_func.sum(Charge.amount), 0))
                .filter(
                    Charge.property_id.in_(property_ids or [-1]),
                    Charge.period_start >= period_start,
                )
                .scalar()
                or 0
            )
            fees = round(collected * 0.08, 2)  # hypothèse : 8 % à documenter côté mandat
            rows.append(
                {
                    "proprietaire": owner.company_name or f"{owner.first_name} {owner.last_name}",
                    "biens": len(properties),
                    "encaisse": round(collected, 2),
                    "charges": round(charges_total, 2),
                    "honoraires_estimes": fees,
                    "net_avant_impots": round(collected - charges_total - fees, 2),
                }
            )
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "period": {"start": period_start.isoformat(), "end": period_end.isoformat()},
            "columns": ["proprietaire", "biens", "encaisse", "charges", "honoraires_estimes", "net_avant_impots"],
            "rows": rows,
            "totals": {"encaisse": round(sum(r["encaisse"] for r in rows), 2)},
        }

    if report_key == "bilan_bien":
        rows = []
        for prop in db.query(Property).all():
            lease = _active_lease_for(db, prop.id)
            collected = (
                db.query(sa_func.coalesce(sa_func.sum(RentPayment.amount_paid), 0))
                .join(Lease, RentPayment.lease_id == Lease.id)
                .filter(Lease.property_id == prop.id)
                .scalar()
                or 0
            )
            charges_total = (
                db.query(sa_func.coalesce(sa_func.sum(Charge.amount), 0))
                .filter(Charge.property_id == prop.id)
                .scalar()
                or 0
            )
            tickets_cost = (
                db.query(sa_func.coalesce(sa_func.sum(MaintenanceTicket.final_cost), 0))
                .filter(MaintenanceTicket.property_id == prop.id)
                .scalar()
                or 0
            )
            rows.append(
                {
                    "bien": f"{prop.reference} — {prop.title}",
                    "ville": prop.city,
                    "occupe": lease is not None,
                    "loyer_mensuel": prop.rent_price,
                    "encaisse_total": round(collected, 2),
                    "charges": round(charges_total, 2),
                    "travaux": round(tickets_cost, 2),
                    "resultat": round(collected - charges_total - tickets_cost, 2),
                }
            )
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "columns": ["bien", "ville", "occupe", "loyer_mensuel", "encaisse_total", "charges", "travaux", "resultat"],
            "rows": rows,
        }

    if report_key == "rapport_activite_agence":
        kpis = dashboard_kpis(db)
        from app.services.crm_service import agent_performance

        performance = agent_performance(
            db,
            date.fromisoformat(params["date_from"]) if params.get("date_from") else None,
            date.fromisoformat(params["date_to"]) if params.get("date_to") else None,
        )
        listings_count = db.query(Listing).count()
        views_total = 0
        for listing in db.query(Listing).all():
            from app.services.crm_service import listing_stats

            views_total += listing_stats(db, listing.id)["totals"]["views"]
        return {
            "title": PREDEFINED_REPORTS[report_key],
            "kpis": kpis,
            "performance_commerciale": performance["global"],
            "annonces": {"count": listings_count, "views_total": views_total},
        }

    if report_key == "rapport_fiscal_annuel":
        rows = []
        payments = (
            db.query(RentPayment, Lease, Property)
            .join(Lease, RentPayment.lease_id == Lease.id)
            .join(Property, Lease.property_id == Property.id)
            .filter(
                RentPayment.due_date >= date(year, 1, 1),
                RentPayment.due_date <= date(year, 12, 31),
            )
            .all()
        )
        per_property: Dict[str, Dict] = {}
        for payment, lease, prop in payments:
            key = f"{prop.reference}"
            bucket = per_property.setdefault(
                key,
                {"bien": f"{prop.reference} — {prop.title}", "adresse": prop.address, "loyers_percus": 0.0, "charges_recuperables": 0.0},
            )
            bucket["loyers_percus"] += (payment.amount_paid or 0)
            bucket["charges_recuperables"] += 0  # détail via régularisation
        for bucket in per_property.values():
            bucket["loyers_percus"] = round(bucket["loyers_percus"], 2)
            rows.append(bucket)
        charges_recoverables = (
            db.query(sa_func.coalesce(sa_func.sum(Charge.amount), 0))
            .filter(Charge.recoverability == ChargeRecoverability.RECOVERABLE)
            .scalar()
            or 0
        )
        return {
            "title": f"{PREDEFINED_REPORTS[report_key]} {year}",
            "columns": ["bien", "adresse", "loyers_percus", "charges_recuperables"],
            "rows": rows,
            "totals": {
                "loyers_percus": round(sum(r["loyers_percus"] for r in rows), 2),
                "charges_recuperables": round(charges_recoverables, 2),
            },
            "mention": "Document de synthèse à faire valider par le comptable avant transmission fiscale.",
        }

    raise ValueError(f"Rapport inconnu : {report_key}")


# ---------------------------------------------------------------------------
# Rapports personnalisés : datasets, filtres, groupement
# ---------------------------------------------------------------------------
DATASETS = {
    "biens": {
        "model": "Property",
        "label": "Biens",
        "fields": {
            "id": "id", "reference": "reference", "titre": "title", "type": "type",
            "statut": "status", "ville": "city", "code_postal": "postal_code",
            "surface": "living_area", "pieces": "rooms", "chambres": "bedrooms",
            "loyer": "rent_price", "charges": "charges", "prix_vente": "sale_price",
            "cree_le": "created_at",
        },
    },
    "baux": {
        "model": "Lease",
        "label": "Baux",
        "fields": {
            "id": "id", "reference": "reference", "bien_id": "property_id",
            "locataire_id": "tenant_id", "statut": "status", "debut": "start_date",
            "fin": "end_date", "loyer": "monthly_rent", "charges": "monthly_charges",
            "depot": "deposit",
        },
    },
    "loyers": {
        "model": "RentPayment",
        "label": "Loyers et paiements",
        "fields": {
            "id": "id", "reference": "reference", "bail_id": "lease_id",
            "periode": "period", "echeance": "due_date", "du": "amount_due",
            "paye": "amount_paid", "statut": "status", "paye_le": "paid_at",
        },
    },
    "impayes": {
        "model": "LatePayment",
        "label": "Impayés",
        "fields": {
            "id": "id", "reference": "reference", "periode": "period",
            "bien_id": "property_id", "locataire_id": "tenant_id",
            "restant_du": "amount_outstanding", "penalites": "penalty_amount",
            "jours_retard": "overdue_days", "etape": "stage", "statut": "status",
        },
    },
    "tickets": {
        "model": "MaintenanceTicket",
        "label": "Tickets maintenance",
        "fields": {
            "id": "id", "reference": "reference", "titre": "title",
            "categorie": "category", "urgence": "urgency", "statut": "status",
            "bien_id": "property_id", "cout_estime": "estimated_cost",
            "cout_final": "final_cost", "ouvert_le": "reported_at",
        },
    },
    "charges": {
        "model": "Charge",
        "label": "Charges",
        "fields": {
            "id": "id", "reference": "reference", "bien_id": "property_id",
            "nature": "charge_type", "categorie": "category", "montant": "amount",
            "recuperable": "recoverability", "debut": "period_start", "fin": "period_end",
        },
    },
    "prospects": {
        "model": "Prospect",
        "label": "Prospects CRM",
        "fields": {
            "id": "id", "reference": "reference", "nom": "last_name",
            "prenom": "first_name", "email": "email", "type": "prospect_type",
            "source": "source", "statut": "status", "budget_min": "budget_min",
            "budget_max": "budget_max", "score": "quality_score",
            "agent": "assigned_agent", "cree_le": "created_at",
        },
    },
    "dossiers": {
        "model": "PipelineDeal",
        "label": "Dossiers commerciaux",
        "fields": {
            "id": "id", "reference": "reference", "titre": "title",
            "prospect_id": "prospect_id", "bien_id": "property_id",
            "etape_id": "stage_id", "statut": "status", "valeur": "estimated_value",
            "commission": "expected_commission", "agent": "assigned_agent",
            "cree_le": "created_at",
        },
    },
    "visites": {
        "model": "Visit",
        "label": "Visites",
        "fields": {
            "id": "id", "reference": "reference", "bien_id": "property_id",
            "prospect_id": "prospect_id", "date": "scheduled_date",
            "debut": "start_time", "fin": "end_time", "statut": "status",
            "agent": "assigned_agent",
        },
    },
    "annonces": {
        "model": "Listing",
        "label": "Annonces",
        "fields": {
            "id": "id", "reference": "reference", "bien_id": "property_id",
            "titre": "title", "prix": "price", "type": "listing_type",
            "statut": "status", "publiee_le": "published_at",
        },
    },
}

MODEL_CLASSES = {
    "Property": Property,
    "Lease": Lease,
    "RentPayment": RentPayment,
    "LatePayment": LatePayment,
    "MaintenanceTicket": MaintenanceTicket,
    "Charge": Charge,
    "Prospect": Prospect,
    "PipelineDeal": PipelineDeal,
    "Visit": Visit,
    "Listing": Listing,
}


def dataset_catalog() -> List[Dict]:
    return [
        {"dataset": key, "label": spec["label"], "fields": sorted(spec["fields"].keys())}
        for key, spec in DATASETS.items()
    ]


def _resolve_column(dataset: str, alias: str) -> SaColumn:
    spec = DATASETS.get(dataset)
    if not spec:
        raise ValueError(f"Dataset inconnu : {dataset}")
    attribute = spec["fields"].get(alias) or (
        alias if alias in spec["fields"].values() else None
    )
    if not attribute:
        raise ValueError(f"Champ inconnu « {alias} » pour le dataset {dataset}")
    model = MODEL_CLASSES[spec["model"]]
    column = getattr(model, attribute, None)
    if column is None:
        raise ValueError(f"Champ non disponible : {alias}")
    return column


def _coerce_value(column, value):
    """Coerce la valeur d'un filtre vers le type de la colonne (enum, date,
    datetime) afin de comparer exactement ce qui est stocké."""
    if value is None:
        return None
    column_type = getattr(column, "type", None)
    if column_type is None:
        return value
    # Colonnes enum : accepter la valeur ou le nom du membre
    enum_class = getattr(column_type, "enum_class", None)
    if enum_class is not None and isinstance(value, str):
        try:
            return enum_class(value)
        except ValueError:
            try:
                return enum_class[value]
            except KeyError:
                return value
    python_type = getattr(column_type, "python_type", None)
    try:
        if python_type is date and isinstance(value, str):
            return date.fromisoformat(value)
        if python_type is datetime and isinstance(value, str):
            return datetime.fromisoformat(value)
    except NotImplementedError:
        pass
    return value


def _apply_filter(query, dataset: str, filter_def: Dict):
    column = _resolve_column(dataset, filter_def["field"])
    operator = filter_def.get("operator", "eq")
    value = _coerce_value(column, filter_def.get("value"))
    second_value = _coerce_value(column, filter_def.get("second_value"))
    if operator == "eq":
        return query.filter(column == value)
    if operator == "ne":
        return query.filter(column != value)
    if operator == "gt":
        return query.filter(column > value)
    if operator == "gte":
        return query.filter(column >= value)
    if operator == "lt":
        return query.filter(column < value)
    if operator == "lte":
        return query.filter(column <= value)
    if operator == "like":
        return query.filter(column.ilike(f"%{value}%"))
    if operator == "in":
        return query.filter(column.in_(value or []))
    if operator == "between":
        return query.filter(column.between(value, second_value))
    raise ValueError(f"Opérateur inconnu : {operator}")


def _serialize(value) -> Any:
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    if hasattr(value, "value"):  # Enum SQLAlchemy
        return value.value
    return value


def run_custom_report(db: Session, report: CustomReport) -> Dict:
    spec = DATASETS.get(report.dataset)
    if not spec:
        raise ValueError(f"Dataset inconnu : {report.dataset}")
    model = MODEL_CLASSES[spec["model"]]
    query = db.query(model)

    for filter_def in report.filters or []:
        query = _apply_filter(query, report.dataset, filter_def)

    aliases = report.fields or list(spec["fields"].keys())
    columns = {alias: _resolve_column(report.dataset, alias) for alias in aliases}

    rows: List[Dict] = []
    grouped = bool(report.group_by)
    if grouped:
        group_aliases = report.group_by
        group_columns = {alias: _resolve_column(report.dataset, alias) for alias in group_aliases}
        numeric_columns = {
            alias: col for alias, col in columns.items()
            if alias not in group_columns and _is_numeric(col)
        }
        selections = list(group_columns.values())
        for alias, col in numeric_columns.items():
            selections.extend([
                sa_func.count(col).label(f"{alias}__count"),
                sa_func.coalesce(sa_func.sum(col), 0).label(f"{alias}__sum"),
                sa_func.coalesce(sa_func.avg(col), 0).label(f"{alias}__avg"),
                sa_func.coalesce(sa_func.min(col), 0).label(f"{alias}__min"),
                sa_func.coalesce(sa_func.max(col), 0).label(f"{alias}__max"),
            ])
        selections.append(sa_func.count(model.id).label("__count"))
        results = query.with_entities(*selections).group_by(*group_columns.values()).all()
        for result in results:
            row: Dict[str, Any] = {}
            for alias in group_aliases:
                row[alias] = _serialize(getattr(result, _attr_name(group_columns[alias])))
            for alias in numeric_columns:
                row[alias] = {
                    "count": getattr(result, f"{alias}__count"),
                    "sum": round(_num(getattr(result, f"{alias}__sum")), 2),
                    "avg": round(_num(getattr(result, f"{alias}__avg")), 2),
                    "min": round(_num(getattr(result, f"{alias}__min")), 2),
                    "max": round(_num(getattr(result, f"{alias}__max")), 2),
                }
            row["count"] = getattr(result, "__count")
            rows.append(row)
        columns_out = list(group_aliases) + [a for a in numeric_columns] + ["count"]
    else:
        entities = query
        ordering = []
        for sort_def in report.sort_by or []:
            col = _resolve_column(report.dataset, sort_def.get("field"))
            direction = sort_def.get("dir", "asc")
            ordering.append(col.desc() if direction == "desc" else col)
        if ordering:
            entities = entities.order_by(*ordering)
        if report.limit:
            entities = entities.limit(report.limit)
        records = entities.all()
        for record in records:
            row = {}
            for alias, col in columns.items():
                row[alias] = _serialize(getattr(record, _attr_name(col)))
            rows.append(row)
        columns_out = aliases

    return {
        "report_id": report.id,
        "name": report.name,
        "dataset": report.dataset,
        "columns": columns_out,
        "grouped": grouped,
        "rows": rows,
        "row_count": len(rows),
        "generated_at": _now().isoformat(),
    }


def _attr_name(column) -> str:
    return column.key if hasattr(column, "key") else str(column).split(".")[-1]


def _is_numeric(column) -> bool:
    try:
        python_type = column.type.python_type
    except NotImplementedError:
        return False
    return python_type in (int, float)


def _num(value) -> float:
    return float(value or 0)


def create_custom_report(db: Session, data, created_by: Optional[str]) -> CustomReport:
    if data.dataset not in DATASETS:
        raise ValueError(f"Dataset inconnu : {data.dataset}")
    invalid = [f for f in (data.fields or []) if f not in DATASETS[data.dataset]["fields"]]
    if invalid:
        raise ValueError(f"Champs inconnus : {', '.join(invalid)}")
    report = CustomReport(created_by=created_by, **data.model_dump(exclude={"filters"}))
    report.filters = [f.model_dump() for f in data.filters]
    db.add(report)
    db.commit()
    db.refresh(report)
    return report


def update_custom_report(db: Session, report_id: int, data) -> CustomReport:
    report = db.query(CustomReport).filter(CustomReport.id == report_id).first()
    if not report:
        raise ValueError("Rapport non trouvé")
    payload = data.model_dump(exclude_unset=True, exclude={"filters"})
    if "dataset" in payload and payload["dataset"] not in DATASETS:
        raise ValueError(f"Dataset inconnu : {payload['dataset']}")
    for field, value in payload.items():
        setattr(report, field, value)
    if data.filters is not None:
        report.filters = [f.model_dump() for f in data.filters]
    db.commit()
    db.refresh(report)
    return report


def delete_custom_report(db: Session, report_id: int) -> None:
    report = db.query(CustomReport).filter(CustomReport.id == report_id).first()
    if not report:
        raise ValueError("Rapport non trouvé")
    db.delete(report)
    db.commit()


def record_execution(
    db: Session,
    report_kind: str,
    report_key: str,
    report_name: str,
    params: Dict,
    output_format: str,
    row_count: int,
    file_path: Optional[str],
    generated_by: Optional[str],
) -> ReportExecution:
    execution = ReportExecution(
        report_kind=report_kind,
        report_key=str(report_key),
        report_name=report_name,
        params=params,
        output_format=output_format,
        row_count=row_count,
        file_path=file_path,
        generated_by=generated_by,
    )
    db.add(execution)
    db.commit()
    db.refresh(execution)
    return execution


def schedule_custom_report(db: Session, report_id: int, frequency: str, recipients: List[str], fmt: str) -> CustomReport:
    report = db.query(CustomReport).filter(CustomReport.id == report_id).first()
    if not report:
        raise ValueError("Rapport non trouvé")
    report.schedule_frequency = frequency
    report.schedule_recipients = recipients
    report.schedule_format = fmt
    report.last_run_status = None
    now = _now()
    if frequency == ScheduleFrequency.NONE.value:
        report.next_run_at = None
    else:
        deltas = {
            ScheduleFrequency.DAILY.value: timedelta(days=1),
            ScheduleFrequency.WEEKLY.value: timedelta(weeks=1),
            ScheduleFrequency.MONTHLY.value: timedelta(days=30),
            ScheduleFrequency.QUARTERLY.value: timedelta(days=91),
        }
        report.next_run_at = now + deltas[frequency]
    db.commit()
    db.refresh(report)
    return report


def run_due_schedules(db: Session) -> Dict:
    """Exécute les rapports dont l'envoi planifié est échu et journalise
    l'expédition vers les destinataires (prestataire à brancher)."""
    now = _now()
    due = (
        db.query(CustomReport)
        .filter(
            CustomReport.is_active == True,  # noqa: E712
            CustomReport.schedule_frequency != ScheduleFrequency.NONE.value,
            CustomReport.next_run_at <= now,
        )
        .all()
    )
    processed = []
    for report in due:
        try:
            result = run_custom_report(db, report)
            execution = record_execution(
                db,
                "custom",
                report.id,
                report.name,
                {"scheduled": True},
                report.schedule_format or "json",
                result["row_count"],
                None,
                "scheduler",
            )
            report.last_run_at = now
            report.last_run_status = "ok"
            deltas = {
                ScheduleFrequency.DAILY.value: timedelta(days=1),
                ScheduleFrequency.WEEKLY.value: timedelta(weeks=1),
                ScheduleFrequency.MONTHLY.value: timedelta(days=30),
                ScheduleFrequency.QUARTERLY.value: timedelta(days=91),
            }
            report.next_run_at = now + deltas[report.schedule_frequency]
            processed.append(
                {
                    "report_id": report.id,
                    "name": report.name,
                    "execution_id": execution.id,
                    "recipients": report.schedule_recipients,
                    "format": report.schedule_format,
                    "rows": result["row_count"],
                }
            )
        except ValueError as exc:
            report.last_run_at = now
            report.last_run_status = f"erreur : {exc}"
            processed.append({"report_id": report.id, "name": report.name, "error": str(exc)})
    db.commit()
    return {"processed": processed, "count": len(processed)}


def get_shared_report(db: Session, token: str) -> Dict:
    report = db.query(CustomReport).filter(CustomReport.share_token == token).first()
    if not report or not report.is_active:
        raise ValueError("Rapport partagé introuvable")
    result = run_custom_report(db, report)
    result["shared"] = True
    return result


# ---------------------------------------------------------------------------
# Exports PDF / Excel / CSV / Word
# ---------------------------------------------------------------------------
def export_pdf(report: Dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Heading1"], fontSize=16,
        textColor=colors.HexColor("#2563eb"), spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontSize=9,
        textColor=colors.grey, spaceAfter=14,
    )
    elements = [Paragraph(report.get("title", "Rapport"), title_style)]
    period = report.get("period")
    if period:
        elements.append(
            Paragraph(
                f"Période : {period['start']} → {period['end']}", subtitle_style
            )
        )
    columns = report.get("columns") or []
    rows = report.get("rows") or []
    if not columns and rows:
        columns = list(rows[0].keys())
    data = [[str(c) for c in columns]]
    for row in rows[:500]:
        data.append([_format_cell(row.get(c)) for c in columns])
    if data.__len__() > 1:
        table = Table(data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(table)
    else:
        elements.append(Paragraph("Aucune donnée sur la période.", subtitle_style))
    totals = report.get("totals")
    if totals:
        elements.append(Spacer(1, 0.6 * cm))
        elements.append(
            Paragraph(
                "Totaux : " + " — ".join(f"{k} = {_format_cell(v)}" for k, v in totals.items()),
                subtitle_style,
            )
        )
    mention = report.get("mention")
    if mention:
        elements.append(Spacer(1, 0.4 * cm))
        elements.append(Paragraph(mention, subtitle_style))
    doc.build(elements)
    return buffer.getvalue()


def export_excel(report: Dict) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = (report.get("title") or "Rapport")[:31]
    columns = report.get("columns") or []
    rows = report.get("rows") or []
    if not columns and rows:
        columns = list(rows[0].keys())
    sheet.append([str(c) for c in columns])
    for cell in sheet[1]:
        cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill("solid", fgColor="2563EB")
    for row in rows:
        sheet.append([_cell_value(row.get(c)) for c in columns])
    for i, column in enumerate(columns, 1):
        length = max([len(str(column))] + [len(str(_cell_value(row.get(column)))) for row in rows[:200]] or [10])
        sheet.column_dimensions[sheet.cell(row=1, column=i).column_letter].width = min(max(length + 2, 10), 45)
    totals = report.get("totals")
    if totals:
        sheet.append([])
        sheet.append([f"{k}: {_cell_value(v)}" for k, v in totals.items()])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def export_csv(report: Dict) -> bytes:
    buffer = io.StringIO()
    writer = csv.writer(buffer, delimiter=";", lineterminator="\n")
    columns = report.get("columns") or []
    rows = report.get("rows") or []
    if not columns and rows:
        columns = list(rows[0].keys())
    writer.writerow(columns)
    for row in rows:
        writer.writerow([_csv_value(row.get(c)) for c in columns])
    return buffer.getvalue().encode("utf-8-sig")


def export_word(report: Dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()
    document.add_heading(report.get("title", "Rapport"), level=1)
    period = report.get("period")
    if period:
        p = document.add_paragraph(f"Période : {period['start']} → {period['end']}")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x67, 0x74, 0x8B)
    columns = report.get("columns") or []
    rows = report.get("rows") or []
    if not columns and rows:
        columns = list(rows[0].keys())
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for i, column in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = str(column)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[:1000]:
        cells = table.add_row().cells
        for i, column in enumerate(columns):
            cells[i].text = _format_cell(row.get(column))
    totals = report.get("totals")
    if totals:
        document.add_paragraph("")
        document.add_paragraph(
            "Totaux : " + " — ".join(f"{k} = {_format_cell(v)}" for k, v in totals.items())
        )
    mention = report.get("mention")
    if mention:
        document.add_paragraph(mention)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _cell_value(value) -> Any:
    if value is None:
        return None
    if isinstance(value, bool):
        return "Oui" if value else "Non"
    if isinstance(value, dict):
        return "; ".join(f"{k}={v}" for k, v in value.items())
    if isinstance(value, (int, float)):
        return value
    return str(value)


def _format_cell(value) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Oui" if value else "Non"
    if isinstance(value, dict):
        return "; ".join(f"{k}={_format_cell(v)}" for k, v in value.items())
    if isinstance(value, float):
        return f"{value:,.2f}".replace(",", " ").replace(".", ",")
    return str(value)


def _csv_value(value) -> Any:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Oui" if value else "Non"
    if isinstance(value, dict):
        return "; ".join(f"{k}={v}" for k, v in value.items())
    return value


EXPORT_BUILDERS = {
    "pdf": export_pdf,
    "excel": export_excel,
    "csv": export_csv,
    "word": export_word,
}


def build_export(report: Dict, fmt: str) -> Tuple[bytes, str, str]:
    """Retourne (contenu, type MIME, extension) pour un format donné."""
    builder = EXPORT_BUILDERS.get(fmt)
    if not builder:
        raise ValueError(f"Format d'export non supporté : {fmt}")
    mime = {
        "pdf": "application/pdf",
        "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv; charset=utf-8",
        "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }[fmt]
    extensions = {"pdf": "pdf", "excel": "xlsx", "csv": "csv", "word": "docx"}
    return builder(report), mime, extensions[fmt]


# ---------------------------------------------------------------------------
# Alertes dashboard
# ---------------------------------------------------------------------------
def _metric_value(db: Session, metric: str) -> float:
    kpis = dashboard_kpis(db)
    mapping = {
        "occupancy_rate": lambda: float(kpis["occupancy_rate_pct"]),
        "monthly_revenue": lambda: float(kpis["revenue"]["month_to_date"]),
        "annual_revenue": lambda: float(kpis["revenue"]["year_to_date"]),
        "unpaid_outstanding": lambda: float(kpis["unpaid"]["outstanding"]),
        "unpaid_count": lambda: float(kpis["unpaid"]["count"]),
        "open_tickets": lambda: float(kpis["maintenance"]["open_tickets"]),
        "leases_expiring_30d": lambda: float(kpis["leases_expiring"]["30"]),
        "mandates_expiring_90d": lambda: float(kpis["mandates_expiring_90d"]),
        "active_prospects": lambda: float(kpis["crm"]["active_prospects"]),
        "open_deals_value": lambda: float(kpis["crm"]["open_deals_value"]),
        "vacant_properties": lambda: float(kpis["properties"]["vacant"]),
    }
    builder = mapping.get(metric)
    if not builder:
        raise ValueError(f"Métrique inconnue : {metric}")
    return builder()


ALERT_METRICS = {
    "occupancy_rate": "Taux d'occupation (%)",
    "monthly_revenue": "Revenus du mois (€)",
    "annual_revenue": "Revenus de l'année (€)",
    "unpaid_outstanding": "Impayés en cours (€)",
    "unpaid_count": "Nombre d'impayés",
    "open_tickets": "Tickets maintenance ouverts",
    "leases_expiring_30d": "Baux expirant sous 30 jours",
    "mandates_expiring_90d": "Mandats expirant sous 90 jours",
    "active_prospects": "Prospects actifs",
    "open_deals_value": "Valeur des dossiers ouverts (€)",
    "vacant_properties": "Biens vacants",
}


def create_alert_rule(db: Session, data) -> AlertRule:
    if data.metric not in ALERT_METRICS:
        raise ValueError(f"Métrique inconnue : {data.metric}")
    rule = AlertRule(**data.model_dump())
    db.add(rule)
    db.commit()
    db.refresh(rule)
    return rule


def update_alert_rule(db: Session, rule_id: int, data) -> AlertRule:
    rule = db.query(AlertRule).filter(AlertRule.id == rule_id).first()
    if not rule:
        raise ValueError("Règle non trouvée")
    payload = data.model_dump(exclude_unset=True)
    if "metric" in payload and payload["metric"] not in ALERT_METRICS:
        raise ValueError(f"Métrique inconnue : {payload['metric']}")
    for field, value in payload.items():
        setattr(rule, field, value)
    db.commit()
    db.refresh(rule)
    return rule


def delete_alert_rule(db: Session, rule_id: int) -> None:
    rule = db.query(AlertRule).filter(AlertRule.id == rule_id).first()
    if not rule:
        raise ValueError("Règle non trouvée")
    db.delete(rule)
    db.commit()


def evaluate_alerts(db: Session) -> Dict:
    """Évalue toutes les règles actives ; déclenche une alerte temps réel si
    le seuil est franchi et que le délai de repos est écoulé."""
    rules = db.query(AlertRule).filter(AlertRule.is_enabled == True).all()  # noqa: E712
    now = _now()
    triggered, checked = [], 0
    for rule in rules:
        try:
            value = _metric_value(db, rule.metric)
        except ValueError:
            continue
        checked += 1
        operators = {
            "<": lambda a, b: a < b,
            "<=": lambda a, b: a <= b,
            ">": lambda a, b: a > b,
            ">=": lambda a, b: a >= b,
            "==": lambda a, b: a == b,
        }
        breached = operators.get(rule.operator, lambda a, b: False)(value, rule.threshold)
        if not breached:
            continue
        if rule.last_triggered_at and rule.cooldown_hours:
            elapsed = (now - _as_utc(rule.last_triggered_at)).total_seconds() / 3600
            if elapsed < rule.cooldown_hours:
                continue
        label = ALERT_METRICS.get(rule.metric, rule.metric)
        message = (
            f"{label} = {value:g} (seuil {rule.operator} {rule.threshold:g} franchi)"
        )
        event = AlertEvent(
            rule_id=rule.id,
            metric=rule.metric,
            value=value,
            threshold=rule.threshold,
            severity=rule.severity,
            message=message,
            channels=rule.channels,
        )
        db.add(event)
        rule.last_triggered_at = now
        triggered.append(
            {
                "rule_id": rule.id,
                "rule": rule.name,
                "severity": rule.severity,
                "value": value,
                "threshold": rule.threshold,
                "message": message,
                "channels": rule.channels,
            }
        )
    db.commit()
    return {"rules_checked": checked, "triggered": triggered, "triggered_count": len(triggered)}


def list_alert_events(db: Session, acknowledged: Optional[bool] = None) -> List[AlertEvent]:
    query = db.query(AlertEvent)
    if acknowledged is True:
        query = query.filter(AlertEvent.acknowledged_at != None)  # noqa: E711
    elif acknowledged is False:
        query = query.filter(AlertEvent.acknowledged_at == None)  # noqa: E711
    return query.order_by(AlertEvent.triggered_at.desc()).limit(200).all()


def acknowledge_alert_event(db: Session, event_id: int, acknowledged: bool) -> AlertEvent:
    event = db.query(AlertEvent).filter(AlertEvent.id == event_id).first()
    if not event:
        raise ValueError("Événement non trouvé")
    event.acknowledged_at = _now() if acknowledged else None
    db.commit()
    db.refresh(event)
    return event
